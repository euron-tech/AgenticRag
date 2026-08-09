"""Seed accounts, departments, and searchable documents for testing.

Accounts are created directly in auth.users/auth.identities because the
service-role key currently in .env is a publishable key and the GoTrue admin
API rejects it. The rows are built to match what GoTrue expects for password
sign-in: a bcrypt hash, a confirmed email, and a matching email identity.

Documents are generated as real .docx/.xlsx/.csv/.md/.json files and pushed
through the project's own loaders, chunker, and embedding client, so retrieval
and citations behave exactly as they will for a genuine upload. The only step
skipped is Supabase Storage, which also needs the secret key.
"""

from __future__ import annotations

import asyncio
import io
import json
import sys
import uuid
from pathlib import Path

ROOT = Path.cwd()
sys.path.insert(0, str(ROOT / "backend"))

import os

for line in (ROOT / ".env").read_text(encoding="utf-8").splitlines():
    if "=" in line and not line.startswith("#"):
        k, _, v = line.partition("=")
        os.environ.setdefault(k.strip(), v.strip())

import asyncpg  # noqa: E402

from app.db import repositories as repo  # noqa: E402
from app.ingestion import loaders  # noqa: E402
from app.ingestion.chunking import chunk_units  # noqa: E402
from app.ingestion.embedding import embed_texts  # noqa: E402

DSN = os.environ["SUPABASE_DB_URL"].strip()
PASSWORD = "Euron@RagDemo2026"

DEPARTMENTS = [
    ("Finance", "finance", "Expense policy, budgets, vendor invoices"),
    ("Human Resources", "human-resources", "Leave policy, handbook, headcount"),
    ("Engineering", "engineering", "Runbooks, API guidelines, service catalog"),
    ("Sales", "sales", "Pricing, territory plans, pipeline"),
]

ACCOUNTS = [
    ("admin@euron.one", "Platform Administrator", "admin", None),
    ("finance@euron.one", "Priya Raman", "user", "finance"),
    ("hr@euron.one", "Arjun Mehta", "user", "human-resources"),
    ("eng@euron.one", "Wei Chen", "user", "engineering"),
    ("sales@euron.one", "Sofia Alvarez", "user", "sales"),
]


# ------------------------------------------------------------- file builders
def docx(title: str, sections: list[tuple[str, list[str]]]) -> bytes:
    import docx as _docx

    document = _docx.Document()
    document.add_heading(title, level=1)
    for heading, paragraphs in sections:
        document.add_heading(heading, level=2)
        for text in paragraphs:
            document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def xlsx(sheets: dict[str, list[list]]) -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    workbook.remove(workbook.active)
    for name, rows in sheets.items():
        sheet = workbook.create_sheet(name[:31])
        for row in rows:
            sheet.append(row)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def text(body: str) -> bytes:
    return body.strip().encode("utf-8")


# ------------------------------------------------------------------ content
def documents() -> dict[str, list[tuple[str, bytes]]]:
    return {
        "finance": [
            ("expense-policy.docx", docx("Global Expense Policy 2026", [
                ("Submission deadlines", [
                    "Expense claims must be submitted within 30 days of the expense being incurred. Claims older than 60 days require written approval from the Finance Director and will not be reimbursed automatically.",
                    "Reimbursement is paid in the next payroll cycle following approval, provided the claim is approved before the 20th of the month.",
                ]),
                ("Travel limits", [
                    "Domestic air travel must be booked in economy class. Business class is permitted only for flights with a scheduled duration exceeding eight hours.",
                    "The nightly hotel cap is INR 8,000 for metro cities and INR 5,500 for all other locations. Amounts above the cap require pre-approval.",
                    "Daily meal allowance is INR 1,500 domestic and USD 60 international. Alcohol is not reimbursable under any circumstance.",
                ]),
                ("Receipts and evidence", [
                    "Any single expense above INR 2,000 requires an itemised receipt. Card statements are not accepted as evidence.",
                    "Lost receipts require a signed declaration countersigned by the employee's line manager.",
                ]),
                ("Approval thresholds", [
                    "Claims up to INR 25,000 are approved by the line manager. Claims between INR 25,000 and INR 100,000 require Finance Manager approval. Anything above INR 100,000 requires CFO sign-off.",
                ]),
            ])),
            ("q3-budget.xlsx", xlsx({"Q3 Budget": [
                ["Cost centre", "Owner", "Budget INR", "Spent INR", "Variance INR", "Status"],
                ["Engineering", "Wei Chen", 12500000, 11840000, 660000, "Under budget"],
                ["Sales", "Sofia Alvarez", 8200000, 8910000, -710000, "Over budget"],
                ["Marketing", "Ravi Kulkarni", 4300000, 3980000, 320000, "Under budget"],
                ["Human Resources", "Arjun Mehta", 2100000, 2050000, 50000, "On track"],
                ["Facilities", "Neha Gupta", 1800000, 1760000, 40000, "On track"],
                ["Customer Support", "Daniel Osei", 3600000, 3720000, -120000, "Over budget"],
            ], "Headcount cost": [
                ["Department", "Headcount", "Average cost INR", "Total INR"],
                ["Engineering", 48, 2400000, 115200000],
                ["Sales", 22, 2100000, 46200000],
                ["Human Resources", 7, 1600000, 11200000],
            ]})),
            ("vendor-invoices.csv", text("""
invoice_id,vendor,category,amount_inr,issued,due,status
INV-2026-0412,Skyline Cloud Services,Infrastructure,1840000,2026-05-02,2026-06-01,Paid
INV-2026-0418,Northwind Consulting,Professional services,960000,2026-05-07,2026-06-06,Paid
INV-2026-0433,Apex Office Supplies,Facilities,128500,2026-05-14,2026-06-13,Overdue
INV-2026-0447,Brightline Recruiting,Recruitment,540000,2026-05-21,2026-06-20,Pending
INV-2026-0455,Delta Security Audit,Compliance,1275000,2026-05-28,2026-06-27,Pending
INV-2026-0461,Orion Data Labs,Data services,432000,2026-06-03,2026-07-03,Paid
INV-2026-0470,Vertex Legal LLP,Legal,880000,2026-06-11,2026-07-11,Disputed
""")),
        ],
        "human-resources": [
            ("leave-policy.docx", docx("Leave and Time Off Policy", [
                ("Annual leave", [
                    "Full-time employees accrue 22 days of paid annual leave per calendar year, accrued monthly at 1.83 days per completed month of service.",
                    "A maximum of 5 unused days may be carried over into the following year. Carried-over days expire on 31 March and are not paid out on exit.",
                ]),
                ("Sick leave", [
                    "Employees are entitled to 12 days of paid sick leave per year. A medical certificate is required for any absence exceeding three consecutive working days.",
                    "Unused sick leave does not carry over and is not encashable.",
                ]),
                ("Parental leave", [
                    "Primary caregivers receive 26 weeks of fully paid parental leave. Secondary caregivers receive 8 weeks, which may be taken in up to three separate blocks within the first year.",
                    "Employees returning from parental leave are entitled to a phased return at 60 percent hours on full pay for the first four weeks.",
                ]),
                ("Notice and approval", [
                    "Annual leave of five days or more requires 14 days notice. Leave requests are approved by the line manager and recorded in the HR system.",
                ]),
            ])),
            ("employee-handbook.md", text("""
# Employee Handbook

## Working hours and flexibility
Core hours are 11:00 to 16:00 local time. Outside core hours employees may
arrange their schedule with their manager. The standard working week is 40 hours.

## Remote and hybrid work
Employees may work remotely up to three days per week. Fully remote arrangements
require director approval and are reviewed every six months. Anyone working
remotely must be reachable during core hours.

## Probation
New joiners serve a probation period of six months. Probation reviews happen at
month three and month six. Notice during probation is two weeks on either side.

## Notice periods
After probation the notice period is 60 days for individual contributors and
90 days for managers and above. Garden leave may be applied at the company's
discretion.

## Performance reviews
Formal reviews run twice a year, in April and October. Ratings are calibrated
across departments before being shared. Compensation changes take effect from
1 July and 1 January respectively.

## Referral bonus
Employees receive INR 75,000 for a successful engineering referral and
INR 45,000 for all other roles, paid after the referred employee completes
six months of service.
""")),
            ("headcount.xlsx", xlsx({"Headcount": [
                ["Department", "Open roles", "Filled", "Attrition %", "Hiring manager"],
                ["Engineering", 9, 48, 8.2, "Wei Chen"],
                ["Sales", 4, 22, 14.6, "Sofia Alvarez"],
                ["Marketing", 2, 11, 9.1, "Ravi Kulkarni"],
                ["Human Resources", 1, 7, 4.3, "Arjun Mehta"],
                ["Customer Support", 6, 19, 21.4, "Daniel Osei"],
                ["Finance", 2, 9, 5.5, "Priya Raman"],
            ]})),
        ],
        "engineering": [
            ("deployment-runbook.md", text("""
# Production Deployment Runbook

## Release windows
Deployments to production are permitted Monday through Thursday between 10:00
and 16:00 IST. No deployments on Friday, on public holidays, or during a
declared incident freeze.

## Pre-deployment checklist
1. CI is green on the release commit.
2. The migration, if any, has been applied to staging and verified.
3. The on-call engineer has acknowledged the release in the deploy channel.
4. Rollback plan is written in the release ticket.

## Blue green procedure
The replacement task set is brought up and held before any production traffic
moves. Smoke tests run against the test listener. Only when those pass is
traffic shifted. The previous version stays running for ten minutes so a
rollback is immediate.

## Rollback
Any CloudWatch alarm firing during the bake window rolls the deployment back
automatically. A manual rollback is performed by stopping the deployment with
auto-rollback enabled. Never roll forward during an incident.

## Severity definitions
SEV1 is a total outage or data loss, with a 15 minute response target.
SEV2 is major degradation affecting many customers, 30 minute response.
SEV3 is minor or cosmetic, handled in the next business day.

## On-call
On-call rotates weekly, handing over on Wednesday at 11:00. The on-call
engineer owns the pager and is not assigned sprint work during their week.
""")),
            ("api-guidelines.docx", docx("Internal API Design Guidelines", [
                ("Versioning", [
                    "All public endpoints are versioned in the path, for example /api/v1. A breaking change requires a new version; existing versions are supported for a minimum of 12 months after deprecation is announced.",
                ]),
                ("Errors", [
                    "Every error response returns a stable machine-readable code, a human-readable message, and the request id. Stack traces are never returned to a client.",
                    "Use 400 for malformed input, 401 for missing or invalid credentials, 403 for a valid identity lacking permission, and 404 when the resource does not exist or is not visible to the caller.",
                ]),
                ("Pagination", [
                    "List endpoints default to 50 items and cap at 200. Pagination is cursor based; offset pagination is not permitted on tables expected to exceed one million rows.",
                ]),
                ("Authentication", [
                    "Services authenticate with short-lived tokens. Long-lived static API keys are prohibited in production. Every endpoint declares its auth requirement explicitly; there is no implicitly public route except the health check.",
                ]),
            ])),
            ("service-catalog.json", text(json.dumps({
                "services": [
                    {"name": "rag-api", "language": "python", "runtime": "fastapi",
                     "owner": "Wei Chen", "tier": 1, "sla": "99.9",
                     "dependencies": ["supabase-postgres", "openai"],
                     "oncall": "platform-rotation"},
                    {"name": "rag-ui", "language": "python", "runtime": "streamlit",
                     "owner": "Wei Chen", "tier": 2, "sla": "99.5",
                     "dependencies": ["rag-api"], "oncall": "platform-rotation"},
                    {"name": "billing-worker", "language": "go", "runtime": "ecs-fargate",
                     "owner": "Priya Raman", "tier": 1, "sla": "99.95",
                     "dependencies": ["postgres", "stripe"], "oncall": "billing-rotation"},
                    {"name": "search-indexer", "language": "python", "runtime": "ecs-fargate",
                     "owner": "Daniel Osei", "tier": 2, "sla": "99.0",
                     "dependencies": ["opensearch", "s3"], "oncall": "platform-rotation"},
                ],
                "escalation": {"sev1": "page-immediately", "sev2": "page-business-hours",
                               "sev3": "ticket"},
            }, indent=2))),
        ],
        "sales": [
            ("pricing-2026.xlsx", xlsx({"Plans": [
                ["Plan", "Monthly USD", "Annual USD", "Seats included", "Extra seat USD", "Support"],
                ["Starter", 49, 490, 5, 12, "Email, 48h"],
                ["Growth", 199, 1990, 25, 9, "Email + chat, 12h"],
                ["Scale", 599, 5990, 100, 7, "Priority, 4h"],
                ["Enterprise", 0, 0, 0, 0, "Dedicated CSM, 1h"],
            ], "Discounts": [
                ["Contract length", "Discount %", "Approval required"],
                ["12 months", 0, "None"],
                ["24 months", 10, "Sales Manager"],
                ["36 months", 18, "VP Sales"],
                ["Non-standard", 25, "CFO and VP Sales"],
            ]})),
            ("territory-plan.docx", docx("FY26 Territory and Quota Plan", [
                ("Territory allocation", [
                    "The India West territory is owned by Sofia Alvarez with an annual quota of USD 1.8 million. India South is owned by Karthik Iyer with a quota of USD 1.4 million.",
                    "EMEA is split into Northern Europe, owned by Lena Fischer at USD 2.2 million, and Southern Europe, owned by Marco Rossi at USD 1.6 million.",
                ]),
                ("Commission structure", [
                    "Commission is 8 percent of closed annual contract value up to quota, and 12 percent on everything above quota. Accelerators reset at the start of each fiscal year.",
                    "Commission is paid in the month following cash collection, not on signature.",
                ]),
                ("Deal desk", [
                    "Any discount above 18 percent, any non-standard payment term, and any contract with a custom liability cap must go through deal desk review before a quote is issued.",
                ]),
            ])),
            ("pipeline.csv", text("""
opportunity,account,stage,value_usd,close_date,owner,probability
OPP-3301,Meridian Logistics,Negotiation,240000,2026-09-15,Sofia Alvarez,70
OPP-3312,Halcyon Retail Group,Proposal,185000,2026-09-30,Karthik Iyer,50
OPP-3325,Northstar Bank,Discovery,410000,2026-11-12,Lena Fischer,20
OPP-3338,Verdant Agritech,Closed Won,96000,2026-07-28,Sofia Alvarez,100
OPP-3344,Cobalt Manufacturing,Negotiation,320000,2026-10-08,Marco Rossi,65
OPP-3350,Lumen Health Systems,Proposal,275000,2026-10-22,Lena Fischer,45
OPP-3361,Pinnacle Media,Closed Lost,150000,2026-08-01,Karthik Iyer,0
""")),
        ],
    }


# --------------------------------------------------------------------- seed
async def main() -> None:
    conn = await asyncpg.connect(DSN, statement_cache_size=0)
    print(f"  password for every seeded account: {PASSWORD}\n")

    # ---- departments
    dept_ids: dict[str, uuid.UUID] = {}
    for name, slug, description in DEPARTMENTS:
        row = await conn.fetchrow(
            """
            insert into departments (name, slug, description)
            values ($1, $2, $3)
            on conflict (slug) do update set name = excluded.name,
                                             description = excluded.description
            returning id
            """,
            name, slug, description,
        )
        dept_ids[slug] = row["id"]
    print(f"  departments: {', '.join(dept_ids)}")

    # ---- accounts
    for email, full_name, role, slug in ACCOUNTS:
        existing = await conn.fetchval("select id from auth.users where email = $1", email)
        if existing:
            user_id = existing
            await conn.execute(
                "update auth.users set encrypted_password = extensions.crypt($2, extensions.gen_salt('bf')) where id = $1",
                user_id, PASSWORD,
            )
        else:
            user_id = uuid.uuid4()
            await conn.execute(
                """
                insert into auth.users (
                    instance_id, id, aud, role, email, encrypted_password,
                    email_confirmed_at, raw_app_meta_data, raw_user_meta_data,
                    created_at, updated_at
                ) values (
                    '00000000-0000-0000-0000-000000000000', $1, 'authenticated',
                    'authenticated', $2,
                    extensions.crypt($3, extensions.gen_salt('bf')),
                    now(),
                    '{"provider":"email","providers":["email"]}'::jsonb,
                    jsonb_build_object('full_name', $4::text),
                    now(), now()
                )
                """,
                user_id, email, PASSWORD, full_name,
            )
            # GoTrue needs a matching email identity or password sign-in fails
            await conn.execute(
                """
                insert into auth.identities (
                    provider_id, user_id, identity_data, provider,
                    last_sign_in_at, created_at, updated_at
                ) values ($1, $2, $3::jsonb, 'email', now(), now(), now())
                """,
                str(user_id), user_id,
                json.dumps({"sub": str(user_id), "email": email,
                            "email_verified": True, "phone_verified": False}),
            )

        await conn.execute(
            """
            insert into profiles (id, email, full_name, role)
            values ($1, $2, $3, $4)
            on conflict (id) do update set role = excluded.role,
                                           full_name = excluded.full_name,
                                           is_active = true
            """,
            user_id, email, full_name, role,
        )
        await conn.execute("delete from user_departments where user_id = $1", user_id)
        targets = list(dept_ids.values()) if role == "admin" else [dept_ids[slug]]
        for dept in targets:
            await conn.execute(
                "insert into user_departments (user_id, department_id) values ($1, $2) "
                "on conflict do nothing",
                user_id, dept,
            )
        scope = "all departments" if role == "admin" else slug
        print(f"    {email:22} {role:6} -> {scope}")

    # GoTrue scans these columns into non-nullable Go strings. A row inserted by
    # SQL leaves them NULL, and every sign-in then fails with the unhelpful
    # "Database error querying schema". The admin API sets them to '' itself.
    token_columns = [
        c for c in (
            "confirmation_token", "recovery_token", "email_change_token_new",
            "email_change", "email_change_token_current", "phone_change",
            "phone_change_token", "reauthentication_token",
        )
        if c in {
            r["column_name"]
            for r in await conn.fetch(
                "select column_name from information_schema.columns "
                "where table_schema='auth' and table_name='users'"
            )
        }
    ]
    await conn.execute(
        "update auth.users set "
        + ", ".join(f"{c} = coalesce({c}, '')" for c in token_columns)
    )
    print(f"    normalised {len(token_columns)} GoTrue token columns (NULL breaks sign-in)")

    # ---- documents
    print()
    total_chunks = 0
    for slug, files in documents().items():
        department_id = dept_ids[slug]
        for filename, data in files:
            result = loaders.load(filename, data)
            chunks = chunk_units(result.units)
            if not chunks:
                print(f"    {slug}/{filename}: produced no chunks, skipped")
                continue
            vectors = await embed_texts([c.content for c in chunks])

            document_id = await conn.fetchval(
                """
                insert into documents (department_id, filename, storage_path, mime_type,
                                       size_bytes, checksum_sha256, status, page_count,
                                       chunk_count, processed_at, metadata)
                values ($1, $2, $3, $4, $5, $6, 'ready', $7, $8, now(),
                        '{"seeded": true}'::jsonb)
                on conflict (department_id, checksum_sha256) do update
                    set status = 'ready', chunk_count = excluded.chunk_count
                returning id
                """,
                department_id, filename, f"seed/{slug}/{filename}",
                loaders.MIME_TYPES.get(loaders.extension_of(filename), "application/octet-stream"),
                len(data),
                __import__("hashlib").sha256(data).hexdigest(),
                result.page_count, len(chunks),
            )
            await repo.delete_chunks(conn, str(document_id))
            await repo.insert_chunks(
                conn,
                document_id=str(document_id),
                department_id=str(department_id),
                chunks=[
                    {"chunk_index": c.chunk_index, "content": c.content,
                     "token_count": c.token_count, "embedding": v, "metadata": c.metadata}
                    for c, v in zip(chunks, vectors)
                ],
            )
            total_chunks += len(chunks)
            print(f"    {slug:16} {filename:26} {len(chunks):3} chunks  ({result.source_type})")

    print(f"\n  total chunks embedded and indexed: {total_chunks}")
    await conn.close()


asyncio.run(main())
