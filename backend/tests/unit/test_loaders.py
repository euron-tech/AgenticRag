from __future__ import annotations

import io
import json

import pytest

from app.core.errors import ValidationFailed
from app.ingestion import loaders


def test_json_flattens_paths_and_groups_by_top_level_key():
    payload = {"invoice": {"number": "INV-42", "total": 100}, "vendor": {"name": "Acme"}}
    result = loaders.load("data.json", json.dumps(payload).encode())
    assert {u.metadata["json_key"] for u in result.units} == {"invoice", "vendor"}
    text = "\n".join(u.text for u in result.units)
    assert "invoice.number: INV-42" in text


def test_malformed_json_fails_with_a_readable_reason():
    with pytest.raises(ValidationFailed) as exc:
        loaders.load("broken.json", b'{"a": ')
    assert "Invalid JSON" in str(exc.value)


def test_csv_repeats_the_header_in_every_unit():
    rows = ["id,name,amount"] + [f"{i},item{i},{i * 10}" for i in range(1, 90)]
    result = loaders.load("data.csv", "\n".join(rows).encode())
    assert len(result.units) > 1
    for unit in result.units:
        assert unit.text.startswith("id | name | amount")
        assert unit.atomic is True


def test_xlsx_records_sheet_and_row_range():
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sales"
    sheet.append(["region", "revenue"])
    sheet.append(["EMEA", 1000])
    buffer = io.BytesIO()
    workbook.save(buffer)

    result = loaders.load("report.xlsx", buffer.getvalue())
    assert result.units
    assert result.units[0].metadata["sheet"] == "Sales"
    assert "Sheet 'Sales' rows" in result.units[0].location


def test_docx_keeps_the_heading_path():
    import docx

    document = docx.Document()
    document.add_heading("Policy", level=1)
    document.add_paragraph("Employees may carry over five days.")
    buffer = io.BytesIO()
    document.save(buffer)

    result = loaders.load("policy.docx", buffer.getvalue())
    assert result.units
    assert result.units[0].metadata["heading_path"] == ["Policy"]


def test_html_strips_script_and_style_content():
    html = b"<html><head><style>.a{color:red}</style></head><body><p>Visible</p>" \
           b"<script>var secret=1</script></body></html>"
    result = loaders.load("page.html", html)
    text = result.units[0].text
    assert "Visible" in text
    assert "secret" not in text and "color:red" not in text


def test_corrupt_pdf_fails_rather_than_returning_nothing():
    with pytest.raises(ValidationFailed):
        loaders.load("broken.pdf", b"%PDF-1.4 this is not really a pdf")


def test_unsupported_extension_names_what_is_supported():
    with pytest.raises(ValidationFailed) as exc:
        loaders.load("archive.zip", b"PK\x03\x04")
    assert ".pdf" in str(exc.value)
